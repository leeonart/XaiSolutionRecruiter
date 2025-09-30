import os
import re
from typing import Dict, List, Tuple, Optional
from docx import Document
import fitz  # PyMuPDF

def extract_text_from_pdf(path: str) -> str:
    """
    Extract text from a PDF file.
    
    Args:
        path: Path to the PDF file
        
    Returns:
        Extracted text as a string
    """
    try:
        txt = ""
        with fitz.open(path) as doc:
            for p in doc:
                txt += p.get_text("text") + "\n"
        return txt
    except Exception as e:
        print(f"Error extracting text from PDF {path}: {e}")
        return f"[Error extracting PDF content: {e}]"

def extract_text_from_docx(path: str) -> str:
    """
    Extract text from a DOCX file with comprehensive content extraction.

    Args:
        path: Path to the DOCX file

    Returns:
        Extracted text as a string
    """
    try:
        doc = Document(path)
        text_parts = []

        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text.strip())

        # Extract from headers
        if hasattr(doc, 'sections'):
            for section in doc.sections:
                if hasattr(section, 'header') and section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(f"Header: {paragraph.text.strip()}")

        # Extract from footers
        if hasattr(doc, 'sections'):
            for section in doc.sections:
                if hasattr(section, 'footer') and section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(f"Footer: {paragraph.text.strip()}")

        # Join all text parts
        full_text = "\n".join(text_parts)

        # Debug logging for empty extractions
        if not full_text.strip():
            print(f"Warning: No text extracted from DOCX {path}")
            print(f"Document has {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
            # Try to get any text at all, even if empty
            all_text = []
            for p in doc.paragraphs:
                all_text.append(f"[Para: '{p.text}']")
            for i, table in enumerate(doc.tables):
                all_text.append(f"[Table {i}: {len(table.rows)} rows]")
            if all_text:
                print(f"Document structure: {' | '.join(all_text[:10])}")
        else:
            print(f"Successfully extracted {len(full_text)} characters from DOCX {os.path.basename(path)}")

        return full_text

    except Exception as e:
        print(f"Error extracting text from DOCX {path}: {e}")
        import traceback
        traceback.print_exc()
        return f"[Error extracting DOCX content: {e}]"

def extract_text_from_txt(path: str) -> str:
    """
    Extract text from a plain text file.
    
    Args:
        path: Path to the text file
        
    Returns:
        File content as a string
    """
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        try:
            # Try with a different encoding if utf-8 fails
            with open(path, 'r', encoding='latin-1') as f:
                return f.read()
        except Exception as e:
            print(f"Error reading text file {path}: {e}")
            return f"[Error reading text file: {e}]"
    except Exception as e:
        print(f"Error reading text file {path}: {e}")
        return f"[Error reading text file: {e}]"

def combine_texts(source_dir: str, output_file: str) -> None:
    """
    Combine text from PDF, DOCX, and TXT files in the source directory.
    
    Args:
        source_dir: Source directory containing files to process
        output_file: Path to the output file
    """
    if not os.path.isdir(source_dir):
        print(f"Error: Source directory '{source_dir}' not found.")
        return
        
    groups: Dict[str, List[Tuple[str, str]]] = {}
    processed_count = 0
    skipped_count = 0
    
    try:
        for root, _, files in os.walk(source_dir):
            for fn in files:
                try:
                    # Extract group ID from filename
                    m = re.search(r'(?<!\d)(\d{4})(?!\d)', fn)
                    if not m:
                        print(f"Skipping {fn} - no group ID found in filename")
                        skipped_count += 1
                        continue
                        
                    grp = m.group(1)
                    file_path = os.path.join(root, fn)
                    
                    # Extract text based on file extension
                    content = ""
                    file_lower = fn.lower()
                    
                    if file_lower.endswith(".pdf"):
                        content = extract_text_from_pdf(file_path)
                    elif file_lower.endswith(".docx") or file_lower.endswith(".doc"):
                        content = extract_text_from_docx(file_path)
                    elif file_lower.endswith(".txt"):
                        content = extract_text_from_txt(file_path)
                    else:
                        print(f"Skipping unsupported file type: {fn}")
                        skipped_count += 1
                        continue
                        
                    if content:
                        groups.setdefault(grp, []).append((fn, content))
                        processed_count += 1
                    else:
                        print(f"No content extracted from {fn}")
                        skipped_count += 1
                        
                except Exception as e:
                    print(f"Error processing file {fn}: {e}")
                    skipped_count += 1

        # Write combined text to output file
        if not groups:
            print("No files were processed successfully. Output file not created.")
            return
            
        os.makedirs(os.path.dirname(os.path.abspath(output_file)), exist_ok=True)
        
        with open(output_file, "w", encoding="utf-8") as out:
            for grp in sorted(groups):
                out.write(f"\n\n{'='*10} Group {grp} {'='*10}\n")
                for fn, txt in sorted(groups[grp]):
                    out.write(f"--- {fn} ---\n{txt}\n\n")
                    
        print(f"Combined text written to {output_file}")
        print(f"Processed {processed_count} files, skipped {skipped_count} files")
        
    except Exception as e:
        print(f"Error combining texts: {e}")
